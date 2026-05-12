"""
pipeline/docx_tailor.py
───────────────────────
In-place python-docx rewriter. Preserves runs (font, size, color, bold)
on unchanged paragraphs; replaces text and applies green color to runs
of modified/added bullets.

Strategy:
  1. Walk paragraphs, classify each as section heading or content.
  2. For each role in tailored.experience, find the matching header by
     title/company substring, then sequentially replace bullets that follow.
  3. For Skills, replace the first non-heading paragraph in the Skills bucket.
  4. New (added) bullets are cloned from the previous bullet's XML so they
     inherit the same numbering / list style.
"""

from __future__ import annotations

import shutil
import subprocess
from concurrent.futures import ThreadPoolExecutor
from copy import deepcopy
from pathlib import Path

from .config import console
from .template_render import render_html

GREEN = (0x0a, 0x66, 0x2c)


def _is_heading(para) -> str | None:
    """Return the heading word (lowercased) if the paragraph looks like a
    section heading. Heuristic: Heading-style paragraphs, or short bold
    ALL-CAPS lines with no punctuation."""
    style_name = (para.style.name or "").lower() if para.style else ""
    if style_name.startswith("heading"):
        return para.text.strip().lower() or None
    text = para.text.strip()
    if not text or len(text.split()) > 4 or any(ch in text for ch in ":,;"):
        return None
    if text == text.upper() and any(c.isalpha() for c in text):
        return text.lower()
    if para.runs and all(r.bold for r in para.runs if r.text.strip()):
        return text.lower()
    return None


_HEADING_BUCKETS = {
    "skills": "skills", "technical skills": "skills", "core competencies": "skills",
    "experience": "experience", "work experience": "experience",
    "professional experience": "experience", "research experience": "experience",
    "education": "education", "projects": "projects", "awards": "awards",
    "certifications": "certifications", "publications": "publications",
    "activities": "activities", "leadership": "leadership", "volunteer": "volunteer",
    "coursework": "coursework", "languages": "languages",
}


def _bucket_for(heading: str) -> str | None:
    return _HEADING_BUCKETS.get(heading.lower())


def _replace_paragraph_text(para, new_text: str, color_green: bool) -> None:
    """Replace text content while preserving the leading run's style.
    Wipes runs after the first; sets new_text on run[0]."""
    if not para.runs:
        run = para.add_run(new_text)
    else:
        first = para.runs[0]
        for r in para.runs[1:]:
            r.text = ""
        first.text = new_text
        run = first
    if color_green:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(*GREEN)


def _clone_bullet_paragraph(template_para, new_text: str, color_green: bool):
    """Clone a List Bullet paragraph by inserting one with the same numbering style.
    Returns the new Paragraph object so callers can chain another insertion."""
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    new_p = deepcopy(template_para._p)
    template_para._p.addnext(new_p)
    # Wipe inline text runs but keep paragraph properties (style, numPr).
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    p_obj = Paragraph(new_p, template_para._parent)
    run = p_obj.add_run(new_text)
    if color_green:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(*GREEN)
    return p_obj


def _convert_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    """Try docx2pdf (Win/macOS) → libreoffice (Linux/Pi) → fail."""
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            return True
    except Exception as e:
        console.print(f"  [yellow]docx2pdf unavailable: {type(e).__name__}: {e}[/yellow]")
    libre = shutil.which("libreoffice") or shutil.which("soffice")
    if libre:
        try:
            subprocess.run(
                [libre, "--headless", "--convert-to", "pdf", "--outdir",
                 str(pdf_path.parent), str(docx_path)],
                check=True, timeout=90, capture_output=True,
            )
            produced = pdf_path.parent / (docx_path.stem + ".pdf")
            if produced.exists() and produced != pdf_path:
                shutil.move(str(produced), str(pdf_path))
            return pdf_path.exists()
        except Exception as e:
            console.print(f"  [yellow]LibreOffice convert failed: {e}[/yellow]")
    return False


def _apply_tailoring(doc, tailored: dict, *, clean: bool) -> None:
    """Mutate ``doc`` in place per ``tailored``. When ``clean`` is True the
    GREEN run colorization is suppressed on added/modified bullets, so the
    resulting docx (and its converted PDF) reads as all-black body text."""
    # ── Section map: list[bucket-or-None per paragraph index] ───────────────
    section_map: list[str | None] = []
    current: str | None = None
    for p in doc.paragraphs:
        h = _is_heading(p)
        if h:
            current = _bucket_for(h)
        section_map.append(current)

    # ── Skills replacement ──────────────────────────────────────────────────
    cats = tailored.get("skills") or []
    skill_text_parts: list[tuple[str, bool]] = []
    for cat in cats:
        for it in cat.get("items") or []:
            skill_text_parts.append((
                it.get("text") or "",
                (it.get("diff") in ("added", "modified")),
            ))
    skills_replaced = False
    if skill_text_parts:
        for idx, p in enumerate(doc.paragraphs):
            if section_map[idx] != "skills":
                continue
            if _is_heading(p):
                continue
            if not skills_replaced:
                flat = ", ".join(t for t, _ in skill_text_parts)
                _replace_paragraph_text(p, flat, color_green=False)
                skills_replaced = True
                if not clean and any(g for _, g in skill_text_parts) and p.runs:
                    from docx.shared import RGBColor
                    p.runs[0].font.color.rgb = RGBColor(*GREEN)
                break

    # ── Experience bullets ──────────────────────────────────────────────────
    roles = tailored.get("experience") or []
    paragraphs = list(doc.paragraphs)
    role_idx = 0
    i = 0
    while i < len(paragraphs) and role_idx < len(roles):
        p = paragraphs[i]
        text = p.text.strip()
        role = roles[role_idx]
        title = (role.get("title") or "").lower()
        company = (role.get("company") or "").lower()
        is_header = (
            (title and title in text.lower())
            or (company and company in text.lower())
        )
        if is_header:
            bullets_to_apply = list(role.get("bullets") or [])
            j = i + 1
            template_p = None
            while j < len(paragraphs):
                pj = paragraphs[j]
                if pj.style and "Bullet" in (pj.style.name or ""):
                    template_p = pj
                    if not bullets_to_apply:
                        _replace_paragraph_text(pj, "", color_green=False)
                    else:
                        b = bullets_to_apply.pop(0)
                        is_changed = b.get("diff") in ("added", "modified")
                        _replace_paragraph_text(
                            pj,
                            b.get("text") or "",
                            color_green=(is_changed and not clean),
                        )
                    j += 1
                else:
                    break
            if bullets_to_apply and template_p is not None:
                anchor = template_p
                for b in bullets_to_apply:
                    anchor = _clone_bullet_paragraph(
                        anchor, b.get("text") or "", color_green=(not clean),
                    )
            i = j
            role_idx += 1
            continue
        i += 1


def tailor_docx_in_place(
    tailored: dict, source_path: Path, base: str, out_dir: Path,
    format_profile: dict | None = None,
) -> dict:
    """Open the source .docx, edit text in-place, save under {base}.docx.

    Outputs:
      {base}.docx           — diff-colored editable artifact
      {base}_final.docx     — clean (all-black) editable artifact
      {base}_preview.html   — HTML preview (template-lib render, used by SPA)
      {base}.pdf            — diff PDF when docx2pdf or LibreOffice available
      {base}_final.pdf      — clean PDF for sending to employers
    """
    from docx import Document

    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    # ── Build both .docx variants in-memory; PDF conversion runs in parallel.
    # Each Document() open + edit + save is fast (<1s); the cost is the
    # docx2pdf / LibreOffice subprocess in _convert_to_pdf (10+s on real
    # resumes). Running the two PDF conversions concurrently halves the
    # perceived tailor latency for .docx uploads.
    docx_path = out_dir / (base + ".docx")
    diff_doc = Document(str(source_path))
    _apply_tailoring(diff_doc, tailored, clean=False)
    diff_doc.save(str(docx_path))
    console.print(f"  [cyan]DOCX saved (in-place) -> {docx_path.name}[/cyan]")

    final_docx_path = out_dir / (base + "_final.docx")
    clean_doc = Document(str(source_path))
    _apply_tailoring(clean_doc, tailored, clean=True)
    clean_doc.save(str(final_docx_path))

    pdf_path = out_dir / (base + ".pdf")
    final_pdf_path = out_dir / (base + "_final.pdf")
    with ThreadPoolExecutor(max_workers=2) as pool:
        pdf_fut = pool.submit(_convert_to_pdf, docx_path, pdf_path)
        final_fut = pool.submit(_convert_to_pdf, final_docx_path, final_pdf_path)
        pdf_ok = pdf_fut.result()
        final_pdf_ok = final_fut.result()
    if final_pdf_ok:
        console.print(f"  [green]Clean DOCX PDF saved -> {final_pdf_path.name}[/green]")

    # ── HTML preview (template-lib render, drives the SPA iframe) ──────────
    html = render_html(tailored, "single_column_modern", format_profile=format_profile)
    html_path = out_dir / (base + "_preview.html")
    html_path.write_text(html, encoding="utf-8")

    return {
        "tex": None,
        "pdf": pdf_path.name if pdf_ok else None,
        "pdf_final": final_pdf_path.name if final_pdf_ok else None,
        "docx": docx_path.name,
        "docx_final": final_docx_path.name,
        "html_preview": html_path.name,
        "base": base,
        "template_id": "in_place_docx",
        "template_confidence": 1.0,
    }
