"""
pipeline/resume.py
──────────────────
Resume I/O helpers: building the demo resume, reading any supported format,
and saving tailored output (plain-text or LaTeX/PDF).

Key API change vs. the old agent.py monolith:
  _read_resume(path) → (text: str, latex_source: str | None)
The caller is responsible for storing the latex_source; there is no global.
"""

import re
from pathlib import Path

from .config import console, OWNER_NAME, OUTPUT_DIR
from .latex import detect_latex, latex_to_plaintext, apply_tailoring_to_latex, compile_latex_to_pdf


# ── Demo resume ────────────────────────────────────────────────────────────────

def _build_demo_resume() -> str:
    """Return a fully generic demo resume with no real personal data.

    This is only used when the user runs the agent without providing a real
    resume file (i.e., demo mode).  Fill in the CLAUDE.md template to replace
    this with your own details.
    """
    return """\
[Your Full Name]
Email: your.email@university.edu
LinkedIn: www.linkedin.com/in/your-profile
Location: [City, State]

OBJECTIVE
[Degree] student seeking a summer internship in [target field, e.g. IC design,
photonics, hardware engineering, FPGA development].

EDUCATION
[University Name] | B.S. [Your Major] | Expected [Graduation Year]
GPA: [X.XX] / 4.00

TECHNICAL SKILLS
[List your tools and software, e.g. MATLAB, Python, Verilog, SPICE]
[List your lab / fab skills, e.g. Photolithography, Cleanroom Processes, SEM]
[List your CAD tools, e.g. SolidWorks, Fusion 360, AutoCAD]
[List other skills, e.g. FPGA, PCB Design, Linux]

PROJECTS
[Project Name 1]
  [Brief description of what you built or researched.]
  [Tools and techniques used.]

[Project Name 2]
  [Brief description.]
  [Outcome or result.]

EXPERIENCE
[Job Title] | [Company / Lab] | [Start Year]–[End Year or Present]
  [What you did — use action verbs and quantify where possible.]
  [Technology or skill applied.]

INTERESTS
[Interest 1] · [Interest 2] · [Interest 3]
"""


# ── PDF text extraction (multi-library fallback chain) ─────────────────────────

# Common Unicode ligatures / typographic replacements produced by PDF fonts.
_LIGATURE_MAP = str.maketrans({
    "\ufb01": "fi",   # ﬁ
    "\ufb02": "fl",   # ﬂ
    "\ufb00": "ff",   # ﬀ
    "\ufb03": "ffi",  # ﬃ
    "\ufb04": "ffl",  # ﬄ
    "\u2019": "'",    # right single quote → apostrophe
    "\u2018": "'",    # left single quote
    "\u201c": '"',    # left double quote
    "\u201d": '"',    # right double quote
    "\u2013": "-",    # en-dash
    "\u2014": "-",    # em-dash
    "\u00a0": " ",    # non-breaking space
    "\u200b": "",     # zero-width space
})


def _normalise_pdf_text(raw: str) -> str:
    """Clean common PDF extraction artifacts:
    - Replace ligatures and curly quotes.
    - Collapse runs of 3+ blank lines to one blank line.
    - Remove carriage returns.
    - Strip trailing whitespace per line.
    - Re-join lines that were broken mid-word by the PDF renderer
      (indicated by no space before a lower-case continuation).
    """
    text = raw.translate(_LIGATURE_MAP)
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Collapse 3+ consecutive blank lines → double blank line.
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Strip trailing spaces on each line.
    lines = [l.rstrip() for l in text.splitlines()]

    # Re-join lines broken mid-word: if the current line ends with a
    # letter and the next starts with a lower-case letter, merge them
    # with a space. Only do this within non-blank runs.
    merged: list[str] = []
    for i, line in enumerate(lines):
        if (merged
                and merged[-1]
                and line
                and merged[-1][-1].isalpha()
                and line[0].islower()):
            merged[-1] = merged[-1] + " " + line
        else:
            merged.append(line)

    return "\n".join(merged)


def _extract_pdf_text(path: Path) -> tuple[str, str]:
    """Try four PDF libraries in order; return (text, library_name).

    Fallback chain:
      1. pypdfium2   — best layout & performance; recommended.
      2. pdfplumber  — fallback layout awareness.
      3. pypdf       — lightweight; handles encrypted-but-readable PDFs.
      4. pdfminer.six — most permissive text extraction.

    Returns ("", "") when all libraries fail or produce empty output.
    """

    # ── 1. pypdfium2 (high performance, good layout) ─────────────────────────
    try:
        import pypdfium2
        pdf = pypdfium2.PdfDocument(str(path))
        try:
            pages = []
            for page in pdf:
                textpage = page.get_textpage()
                pages.append(textpage.get_text_range() or "")
        finally:
            pdf.close()
        text = _normalise_pdf_text("\n".join(pages))
        if text.strip():
            return text, "pypdfium2"
    except ImportError:
        pass
    except Exception as e:
        console.print(f"  [yellow]pypdfium2 failed ({e}) — trying pdfplumber[/yellow]")

    # ── 2. pdfplumber ────────────────────────────────────────────────────────
    try:
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            pages: list[str] = []
            for page in pdf.pages:
                # extract_text(x_tolerance, y_tolerance) reduces spurious
                # merged words in tightly typeset resumes.
                t = page.extract_text(x_tolerance=2, y_tolerance=4) or ""
                pages.append(t)
        text = _normalise_pdf_text("\n".join(pages))
        if text.strip():
            return text, "pdfplumber"
    except ImportError:
        console.print(
            "  [dim]pdfplumber not installed — trying pypdf "
            "(pip install pdfplumber for best results)[/dim]"
        )
    except Exception as e:
        console.print(f"  [yellow]pdfplumber failed ({e}) — trying pypdf[/yellow]")

    # ── 3. pypdf ─────────────────────────────────────────────────────────────
    try:
        import pypdf  # pypdf >= 3.x (successor to PyPDF2)
        with open(str(path), "rb") as f:
            reader = pypdf.PdfReader(f)
            pages = [page.extract_text() or "" for page in reader.pages]
        text = _normalise_pdf_text("\n".join(pages))
        if text.strip():
            return text, "pypdf"
    except ImportError:
        # Try the legacy PyPDF2 name before giving up on this tier.
        try:
            import PyPDF2
            with open(str(path), "rb") as f:
                reader = PyPDF2.PdfReader(f)
                pages = [page.extract_text() or "" for page in reader.pages]
            text = _normalise_pdf_text("\n".join(pages))
            if text.strip():
                return text, "PyPDF2"
        except ImportError:
            console.print(
                "  [dim]pypdf not installed — trying pdfminer "
                "(pip install pypdf for a better fallback)[/dim]"
            )
        except Exception as e:
            console.print(f"  [yellow]pypdf failed ({e}) — trying pdfminer[/yellow]")
    except Exception as e:
        console.print(f"  [yellow]pypdf failed ({e}) — trying pdfminer[/yellow]")

    # ── 4. pdfminer.six ───────────────────────────────────────────────────────
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        raw = pdfminer_extract(str(path))
        text = _normalise_pdf_text(raw or "")
        if text.strip():
            return text, "pdfminer.six"
    except ImportError:
        console.print(
            "  [yellow]No PDF library available. Install at least one of: "
            "pdfplumber, pypdf, pdfminer.six[/yellow]"
        )
    except Exception as e:
        console.print(f"  [yellow]pdfminer failed ({e})[/yellow]")

    return "", ""


# ── Resume reader ──────────────────────────────────────────────────────────────

def _read_docx(path: Path) -> str:
    """Extract text from a .docx, including paragraphs AND table cells.

    Many resumes use tables (sidebar layouts, two-column templates) where
    the candidate's name, contact info, or experience lives entirely inside
    table cells — paragraph-only extraction misses everything.
    """
    from docx import Document
    doc = Document(str(path))
    chunks: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            chunks.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = "\n".join(
                    p.text for p in cell.paragraphs if p.text.strip()
                ).strip()
                if cell_text:
                    chunks.append(cell_text)
    return "\n".join(chunks)


def _read_rtf(path: Path) -> str:
    """Best-effort RTF extraction. Tries striprtf if installed, otherwise
    falls back to a minimal regex strip — good enough for plain RTF resumes.
    """
    raw = path.read_text(encoding="utf-8", errors="replace")
    try:
        from striprtf.striprtf import rtf_to_text  # type: ignore
        return rtf_to_text(raw)
    except ImportError:
        # Minimal RTF strip: drop control words, control symbols, and braces.
        # Not perfect for complex documents but recovers the body of a
        # straightforward résumé without a new dependency.
        text = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", raw)
        text = re.sub(r"\\[^a-zA-Z]", "", text)
        text = re.sub(r"[{}]", "", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


def _read_resume(path: Path) -> tuple[str, str | None]:
    """Read *path* and return ``(plaintext, latex_source_or_None)``.

    If the file is LaTeX (or contains LaTeX markup), *latex_source* holds the
    original source so callers can later produce a compiled PDF output.
    Returns ("", None) if the file isn't a supported format or if extraction
    yields no text — the caller decides how to surface that to the user.
    """
    if not path.exists():
        console.print(f"  [yellow]File not found: {path}[/yellow]")
        return "", None

    suffix = path.suffix.lower()

    if suffix == ".tex":
        raw = path.read_text(encoding="utf-8", errors="replace")
        console.print("  [cyan]LaTeX resume detected — converting to plain text for parsing.[/cyan]")
        return latex_to_plaintext(raw), raw

    if suffix in (".txt", ".md"):
        try:
            raw = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            raw = path.read_text(encoding="utf-8", errors="replace")
        if detect_latex(raw):
            console.print(
                "  [cyan]LaTeX content detected — converting to plain text for parsing.[/cyan]"
            )
            return latex_to_plaintext(raw), raw
        return raw, None

    if suffix == ".pdf":
        text, method = _extract_pdf_text(path)
        if text:
            console.print(f"  [dim]PDF extracted via {method}[/dim]")
            return text, None
        console.print(
            "  [yellow]PDF returned no text via any library. The file is likely "
            "a scanned/image-only PDF — install an OCR pipeline or upload the "
            "source .docx / .tex / .txt instead.[/yellow]"
        )
        return "", None

    if suffix == ".docx":
        try:
            text = _read_docx(path)
            if text.strip():
                return text, None
            console.print(
                "  [yellow]DOCX parsed but produced no text — the document may "
                "be empty or all content may live inside images.[/yellow]"
            )
            return "", None
        except ImportError:
            console.print("  [yellow]python-docx missing — pip install python-docx[/yellow]")
            return "", None
        except Exception as exc:
            console.print(f"  [yellow]DOCX read failed: {exc}[/yellow]")
            return "", None

    if suffix == ".doc":
        # Legacy Word .doc isn't natively supported by python-docx. Try a
        # binary-safe text extract; if that yields nothing, ask the user
        # to convert.
        try:
            raw = path.read_bytes()
            text = re.sub(rb"[^\x09\x0a\x0d\x20-\x7e]+", b" ", raw)
            decoded = text.decode("utf-8", errors="replace").strip()
            decoded = re.sub(r"\s{3,}", "\n", decoded)
            if len(decoded) > 200:
                return decoded, None
        except Exception:
            pass
        console.print(
            "  [yellow]Legacy .doc not natively supported — please save as "
            ".docx, .pdf, or .txt and re-upload.[/yellow]"
        )
        return "", None

    if suffix == ".rtf":
        try:
            text = _read_rtf(path)
            return text, None
        except Exception as exc:
            console.print(f"  [yellow]RTF read failed: {exc}[/yellow]")
            return "", None

    # Generic text fallback (unknown extensions that happen to be UTF-8).
    try:
        raw = path.read_text(encoding="utf-8")
        return raw, None
    except UnicodeDecodeError:
        try:
            raw = path.read_text(encoding="utf-8", errors="replace")
            return raw, None
        except Exception:
            pass
    except Exception:
        pass
    console.print(
        f"  [yellow]Cannot read {path.name} ({suffix}) — supported formats: "
        ".pdf, .docx, .doc, .txt, .md, .tex, .rtf[/yellow]"
    )
    return "", None


# ── Tailored resume output ─────────────────────────────────────────────────────

def _merge_experience(profile_experience: list, tailored_bullets: list) -> list:
    """Overlay LLM-tailored bullets onto the profile's original experience entries.

    Returns a list of ``{title, company, dates, bullets}`` dicts.  For each role
    in *profile_experience*, looks for a matching entry in *tailored_bullets*
    (by case-insensitive substring match against ``role``) and substitutes the
    tailored bullets when present.  Roles with no match keep their original
    bullets, so the rendered resume is always complete.

    If *profile_experience* is empty but *tailored_bullets* is not, falls back
    to the tailored entries directly (best effort).
    """
    if not profile_experience:
        return [
            {"title": e.get("role", ""), "company": "", "dates": "",
             "bullets": e.get("bullets", [])}
            for e in (tailored_bullets or [])
        ]

    tailored_lookup = {}
    for entry in (tailored_bullets or []):
        key = (entry.get("role") or "").strip().lower()
        if key and entry.get("bullets"):
            tailored_lookup[key] = entry["bullets"]

    merged = []
    for role in profile_experience:
        title   = role.get("title", "") or ""
        company = role.get("company", "") or ""
        bullets = role.get("bullets", []) or []
        # Try matching tailored entry by title substring (either direction).
        title_l = title.lower()
        match   = None
        for k, v in tailored_lookup.items():
            if k and (k in title_l or title_l in k):
                match = v
                break
        merged.append({
            "title":   title,
            "company": company,
            "dates":   role.get("dates", "") or "",
            "bullets": match if match else bullets,
        })
    return merged


def _latex_escape(s: str) -> str:
    """Escape LaTeX-special characters in *s*."""
    if not s:
        return ""
    return (
        str(s)
        .replace("\\", r"\textbackslash{}")
        .replace("&",  r"\&")
        .replace("%",  r"\%")
        .replace("$",  r"\$")
        .replace("#",  r"\#")
        .replace("_",  r"\_")
        .replace("{",  r"\{")
        .replace("}",  r"\}")
        .replace("~",  r"\textasciitilde{}")
        .replace("^",  r"\textasciicircum{}")
    )


def _render_resume_latex(profile: dict, tailored: dict, job: dict,
                         resume_text: str = "") -> str:
    """Build a minimal, self-contained LaTeX resume from profile + tailoring."""
    e = _latex_escape

    name     = profile.get("name") or OWNER_NAME
    email    = profile.get("email", "")
    linkedin = profile.get("linkedin", "")
    location = profile.get("location", "")

    skills = tailored.get("skills_reordered") or profile.get("top_hard_skills") or []
    experience = _merge_experience(
        profile.get("experience") or [],
        tailored.get("experience_bullets") or [],
    )
    projects  = profile.get("projects")  or []
    education = profile.get("education") or []

    contact_bits = [b for b in (email, linkedin, location) if b]

    out: list = []
    out.append(r"\documentclass[11pt,letterpaper]{article}")
    out.append(r"\usepackage[margin=0.75in]{geometry}")
    out.append(r"\usepackage{enumitem}")
    out.append(r"\usepackage{titlesec}")
    out.append(r"\setlist[itemize]{leftmargin=*,nosep}")
    out.append(r"\titleformat{\section}{\large\bfseries\uppercase}{}{0pt}{}[\titlerule]")
    out.append(r"\titlespacing*{\section}{0pt}{8pt}{4pt}")
    out.append(r"\pagenumbering{gobble}")
    out.append(r"\begin{document}")
    out.append(r"\begin{center}")
    out.append(r"{\LARGE \textbf{" + e(name) + r"}}\\[2pt]")
    if contact_bits:
        out.append(r"\small " + " \\textbullet{} ".join(e(b) for b in contact_bits))
    out.append(r"\end{center}")

    order = (
        tailored.get("section_order")
        or ["Skills", "Experience", "Projects", "Education"]
    )
    for section in order:
        sec = (section or "").lower()
        if sec == "skills" and skills:
            out.append(r"\section*{Skills}")
            out.append(", ".join(e(s) for s in skills))
        elif sec == "experience" and experience:
            out.append(r"\section*{Experience}")
            for role in experience:
                header = " \\textbar{} ".join(
                    e(b) for b in (role.get("title"), role.get("company"), role.get("dates")) if b
                )
                if header:
                    out.append(r"\noindent\textbf{" + header + r"}\\")
                bullets = role.get("bullets") or []
                if bullets:
                    out.append(r"\begin{itemize}")
                    for b in bullets:
                        out.append(r"\item " + e(b))
                    out.append(r"\end{itemize}")
        elif sec == "projects" and projects:
            out.append(r"\section*{Projects}")
            for p in projects:
                pname = p.get("name", "")
                desc  = p.get("description", "")
                used  = p.get("skills_used") or []
                if pname:
                    out.append(r"\noindent\textbf{" + e(pname) + r"}\\")
                if desc:
                    out.append(e(desc) + r"\\")
                if used:
                    out.append(r"\textit{Skills:} " + e(", ".join(used)))
        elif sec == "education" and education:
            out.append(r"\section*{Education}")
            for ed in education:
                bits = [ed.get("degree", ""), ed.get("institution", ""), ed.get("year", "")]
                line = " \\textbar{} ".join(e(b) for b in bits if b)
                if line:
                    out.append(r"\noindent " + line + r"\\")
                if ed.get("gpa"):
                    out.append(r"GPA: " + e(ed["gpa"]) + r"\\")

    # Fallback: if nothing structured was rendered, embed the raw resume text
    # so the output file is never effectively empty.
    rendered_any = bool(skills or experience or projects or education)
    if not rendered_any and (resume_text or "").strip():
        out.append(r"\section*{Resume}")
        for line in resume_text.splitlines():
            line = line.strip()
            if line:
                out.append(e(line) + r"\\")

    if tailored.get("cover_letter"):
        out.append(r"\vspace{12pt}\hrule\vspace{8pt}")
        out.append(r"\section*{Cover Letter}")
        for para in str(tailored["cover_letter"]).split("\n\n"):
            out.append(e(para.strip()) + r"\\[4pt]")

    out.append(r"\end{document}")
    return "\n".join(out)


def _render_resume_pdf_reportlab(pdf_path: Path, profile: dict,
                                 tailored: dict, job: dict,
                                 resume_text: str = "",
                                 format_profile: dict | None = None) -> bool:
    """Render the same structured resume to a PDF via reportlab.

    Returns True on success, False if reportlab is unavailable or rendering
    fails for any reason.

    *format_profile*, when supplied, lets the caller mirror the source
    PDF's visual aesthetic.  Recognised keys (all optional):

      body_font_size:    int   — drives body / role / contact size
      header_font_size:  int   — drives section heading size
      accent_color:      str   — "#rrggbb" used for headings + rules

    The defaults below match the previous hard-coded values, so PDF
    output is stable when no fingerprint is available (demo / .txt).
    """
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, HRFlowable,
        )
    except ImportError:
        console.print(
            "  [yellow]reportlab missing - pip install reportlab "
            "(skipping PDF generation).[/yellow]"
        )
        return False

    name     = profile.get("name") or OWNER_NAME
    email    = profile.get("email", "")
    linkedin = profile.get("linkedin", "")
    location = profile.get("location", "")

    skills = tailored.get("skills_reordered") or profile.get("top_hard_skills") or []
    experience = _merge_experience(
        profile.get("experience") or [],
        tailored.get("experience_bullets") or [],
    )
    projects  = profile.get("projects")  or []
    education = profile.get("education") or []

    # ── Style overrides from the source PDF's format fingerprint ─────────
    fp = format_profile or {}
    body_size   = float(fp.get("body_font_size")   or 10)
    header_size = float(fp.get("header_font_size") or 11)
    # Sanity-clamp: a malformed fingerprint from a corrupt PDF could
    # produce 1pt or 200pt headers — cap to ranges that print sanely.
    body_size   = max(8.5, min(12.5, body_size))
    header_size = max(body_size + 0.5, min(14, header_size))
    name_size   = round(min(22, header_size + 6.5))
    accent      = fp.get("accent_color") or "#1F4E79"
    if not (isinstance(accent, str) and len(accent) == 7 and accent.startswith("#")):
        accent = "#1F4E79"

    styles = getSampleStyleSheet()
    h_name    = ParagraphStyle("h_name", parent=styles["Title"],
                               fontSize=name_size,
                               spaceAfter=2, alignment=1)
    h_contact = ParagraphStyle("h_contact", parent=styles["Normal"],
                               fontSize=max(8, body_size - 1.5),
                               alignment=1, spaceAfter=8)
    h_section = ParagraphStyle("h_section", parent=styles["Heading2"],
                               fontSize=header_size,
                               textColor=accent, spaceBefore=8, spaceAfter=2)
    h_role    = ParagraphStyle("h_role", parent=styles["Normal"],
                               fontSize=body_size, leading=body_size + 2,
                               spaceAfter=2)
    h_body    = ParagraphStyle("h_body", parent=styles["Normal"],
                               fontSize=body_size, leading=body_size + 2)

    def _esc(s):
        return (str(s or "")
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;"))

    story: list = []
    story.append(Paragraph(_esc(name), h_name))
    contact_bits = [b for b in (email, linkedin, location) if b]
    if contact_bits:
        story.append(Paragraph(" &bull; ".join(_esc(b) for b in contact_bits), h_contact))

    def _add_section(title: str):
        story.append(Paragraph(_esc(title.upper()), h_section))
        story.append(HRFlowable(width="100%", thickness=0.5, color=accent,
                                spaceBefore=0, spaceAfter=4))

    order = (
        tailored.get("section_order")
        or ["Skills", "Experience", "Projects", "Education"]
    )
    for section in order:
        sec = (section or "").lower()
        if sec == "skills" and skills:
            _add_section("Skills")
            story.append(Paragraph(_esc(", ".join(skills)), h_body))
        elif sec == "experience" and experience:
            _add_section("Experience")
            for role in experience:
                header = " | ".join(b for b in (
                    role.get("title"), role.get("company"), role.get("dates")
                ) if b)
                if header:
                    story.append(Paragraph(f"<b>{_esc(header)}</b>", h_role))
                bullets = role.get("bullets") or []
                if bullets:
                    story.append(ListFlowable(
                        [ListItem(Paragraph(_esc(b), h_body), leftIndent=12)
                         for b in bullets],
                        bulletType="bullet", start="circle", leftIndent=12,
                    ))
                story.append(Spacer(1, 4))
        elif sec == "projects" and projects:
            _add_section("Projects")
            for p in projects:
                pname = p.get("name", "")
                desc  = p.get("description", "")
                used  = p.get("skills_used") or []
                if pname:
                    story.append(Paragraph(f"<b>{_esc(pname)}</b>", h_role))
                if desc:
                    story.append(Paragraph(_esc(desc), h_body))
                if used:
                    story.append(Paragraph(
                        f"<i>Skills:</i> {_esc(', '.join(used))}", h_body
                    ))
                story.append(Spacer(1, 4))
        elif sec == "education" and education:
            _add_section("Education")
            for ed in education:
                bits = [ed.get("degree", ""), ed.get("institution", ""),
                        ed.get("year", "")]
                line = " | ".join(b for b in bits if b)
                if line:
                    story.append(Paragraph(_esc(line), h_body))
                if ed.get("gpa"):
                    story.append(Paragraph(f"GPA: {_esc(ed['gpa'])}", h_body))

    # Fallback: embed raw resume text when nothing structured rendered.
    rendered_any = bool(skills or experience or projects or education)
    if not rendered_any and (resume_text or "").strip():
        _add_section("Resume")
        for para in resume_text.split("\n\n"):
            para = para.strip()
            if para:
                story.append(Paragraph(_esc(para).replace("\n", "<br/>"), h_body))
                story.append(Spacer(1, 4))

    if tailored.get("cover_letter"):
        story.append(Spacer(1, 12))
        story.append(HRFlowable(width="100%", thickness=0.5, color="#888888"))
        _add_section("Cover Letter")
        for para in str(tailored["cover_letter"]).split("\n\n"):
            if para.strip():
                story.append(Paragraph(_esc(para.strip()), h_body))
                story.append(Spacer(1, 4))

    try:
        SimpleDocTemplate(
            str(pdf_path), pagesize=LETTER,
            leftMargin=0.75 * inch, rightMargin=0.75 * inch,
            topMargin=0.6 * inch, bottomMargin=0.6 * inch,
            title=f"{name} Resume", author=name,
        ).build(story)
        return True
    except Exception as e:
        console.print(f"  [yellow]reportlab PDF error: {e}[/yellow]")
        return False


def _save_tailored_resume(job: dict, tailored: dict, profile: dict = None,
                          latex_source: str = None,
                          resume_text: str = "",
                          output_dir: Path = None,
                          owner_name: str = None,
                          format_profile: dict | None = None,
                          source_format: str | None = None,
                          source_bytes_path: Path | None = None) -> dict:
    """Write the tailored resume to OUTPUT_DIR.

    Returns ``{"tex": str|None, "pdf": str|None, "docx": str|None,
              "html_preview": str|None, "base": str, "template_id": str|None,
              "template_confidence": float|None}``.

    Dispatch by source_format:
      tex   → in-place LaTeX rewrite (preserves user's .tex template)
      docx  → in-place python-docx rewrite (preserves runs/styles)
      pdf   → template-library match, render via Jinja2 → WeasyPrint
      else  → default template (single_column_classic)

    Backwards-compatible: when source_format is None or the legacy v1 dict
    is passed, falls through to the template-library path.
    """
    from .tailored_schema import (
        SCHEMA_VERSION, default_v2, legacy_to_v2, validate_v2,
    )
    from .template_match import pick_template
    from .template_render import render_html, render_pdf

    safe = lambda s: re.sub(r"[^a-zA-Z0-9_\-]", "_", s)
    name = owner_name or OWNER_NAME
    base = (
        f"{safe(name)}_Resume_{safe(job.get('company', ''))}"
        f"_{safe(job.get('title', ''))}"
    )
    out_dir = output_dir or OUTPUT_DIR
    out_dir.mkdir(parents=True, exist_ok=True)

    # ── Normalize tailored to v2 ────────────────────────────────────────────
    if not isinstance(tailored, dict):
        tailored = {}
    if tailored.get("schema_version") == SCHEMA_VERSION:
        v2 = validate_v2(tailored) or default_v2(profile)
    else:
        v2 = legacy_to_v2(tailored, profile)
    if owner_name:
        v2["name"] = owner_name

    # ── Source-format dispatch ──────────────────────────────────────────────
    src = (source_format or "").lower()
    if src == "tex" and latex_source:
        from .latex_tailor import tailor_latex_in_place
        return tailor_latex_in_place(
            v2, latex_source=latex_source, base=base, out_dir=out_dir,
            format_profile=format_profile,
        )
    if src == "docx" and source_bytes_path and Path(source_bytes_path).exists():
        from .docx_tailor import tailor_docx_in_place
        return tailor_docx_in_place(
            v2, source_path=Path(source_bytes_path), base=base, out_dir=out_dir,
            format_profile=format_profile,
        )

    # PDF / unknown / default → template library
    template_id, confidence = pick_template(format_profile, resume_text)
    # 1. Diff/preview HTML — green-highlighted, drives the in-page iframe.
    html_diff = render_html(v2, template_id, format_profile=format_profile)
    html_path = out_dir / (base + "_preview.html")
    html_path.write_text(html_diff, encoding="utf-8")
    pdf_path = out_dir / (base + ".pdf")
    pdf_ok = render_pdf(html_diff, pdf_path)
    if pdf_ok:
        console.print(f"  [green]PDF saved -> {pdf_path.name} (template={template_id})[/green]")
    else:
        console.print(f"  [yellow]No PDF backend available — only HTML preview at {html_path.name}.[/yellow]")
    # 2. Clean/final PDF — same template, all-black body. This is the file
    # users actually attach to applications. Best-effort: if WeasyPrint isn't
    # available we leave it None, the diff PDF still works as a fallback.
    final_pdf_path = out_dir / (base + "_final.pdf")
    html_clean = render_html(v2, template_id, format_profile=format_profile, clean=True)
    final_pdf_ok = render_pdf(html_clean, final_pdf_path)
    if final_pdf_ok:
        console.print(f"  [green]Clean PDF saved -> {final_pdf_path.name}[/green]")
    return {
        "tex": None,
        "pdf": pdf_path.name if pdf_ok else None,
        "pdf_final": final_pdf_path.name if final_pdf_ok else None,
        "docx": None,
        "html_preview": html_path.name,
        "base": base,
        "template_id": template_id,
        "template_confidence": round(float(confidence), 2),
    }
