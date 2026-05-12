"""Build deterministic .docx fixture for tests.

Run once: ``python tests/fixtures/build_resumes.py``
The produced files are committed to the repo so the test suite
doesn't need to regenerate them on every run.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

HERE = Path(__file__).parent / "resumes"
HERE.mkdir(exist_ok=True)


def build_modern_sans():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(10)

    head = doc.add_paragraph()
    run = head.add_run("Jane Doe")
    run.bold = True
    run.font.size = Pt(20)
    doc.add_paragraph("jane@example.com  ·  Berkeley, CA")

    h_skills = doc.add_paragraph()
    h_skills.add_run("SKILLS").bold = True
    doc.add_paragraph("Python, C++, Verilog, MATLAB")

    h_exp = doc.add_paragraph()
    h_exp.add_run("EXPERIENCE").bold = True
    p1 = doc.add_paragraph()
    p1.add_run("Intern · Acme Corp · 2024").bold = True
    doc.add_paragraph("Built a thing for the team", style="List Bullet")
    doc.add_paragraph("Tested another thing", style="List Bullet")

    p2 = doc.add_paragraph()
    p2.add_run("Research Assistant · Cal Photonics Lab · 2023").bold = True
    doc.add_paragraph("Aligned an interferometer", style="List Bullet")

    h_edu = doc.add_paragraph()
    h_edu.add_run("EDUCATION").bold = True
    doc.add_paragraph(
        "University of California, Berkeley — B.S. Electrical Engineering — 2025",
    )

    doc.save(str(HERE / "modern_sans.docx"))


if __name__ == "__main__":
    build_modern_sans()
    print("Built fixtures in", HERE)
