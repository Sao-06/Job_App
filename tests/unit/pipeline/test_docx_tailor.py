"""Unit tests for pipeline.docx_tailor."""
from pathlib import Path

import pytest

from pipeline.docx_tailor import tailor_docx_in_place
from tests.fakes import jane_doe_tailored_v2

pytestmark = pytest.mark.unit

FIXTURE = (
    Path(__file__).resolve().parents[2] / "fixtures" / "resumes" / "modern_sans.docx"
)


def _v2():
    """DOCX path needs the added-bullet variant to exercise the clone path."""
    return jane_doe_tailored_v2(with_added_bullet=True)


def test_docx_in_place_replaces_modified_bullet(tmp_path):
    out = tailor_docx_in_place(_v2(), source_path=FIXTURE, base="resume", out_dir=tmp_path)
    assert out["docx"] == "resume.docx"
    docx_path = tmp_path / "resume.docx"
    assert docx_path.exists()
    from docx import Document
    doc = Document(str(docx_path))
    texts = [p.text for p in doc.paragraphs]
    assert any("Verilog testbench" in t for t in texts)
    assert any("Aligned an interferometer" in t for t in texts)


def test_docx_in_place_appends_added_bullet(tmp_path):
    tailor_docx_in_place(_v2(), source_path=FIXTURE, base="resume", out_dir=tmp_path)
    from docx import Document
    doc = Document(str(tmp_path / "resume.docx"))
    texts = [p.text for p in doc.paragraphs]
    assert any("AXI4 transaction generators" in t for t in texts)


def test_docx_in_place_returns_html_preview(tmp_path):
    out = tailor_docx_in_place(_v2(), source_path=FIXTURE, base="resume", out_dir=tmp_path)
    assert out["html_preview"] is not None
    assert (tmp_path / out["html_preview"]).exists()


def test_docx_in_place_template_id_is_in_place_docx(tmp_path):
    out = tailor_docx_in_place(_v2(), source_path=FIXTURE, base="resume", out_dir=tmp_path)
    assert out["template_id"] == "in_place_docx"
    assert out["template_confidence"] == 1.0


def test_docx_in_place_clean_has_no_green_runs(tmp_path):
    """The {base}_final.docx (clean variant) must NOT have green-colored
    runs — that file is what gets converted to the all-black PDF the user
    attaches to applications. Diff highlights belong only in the in-page
    preview iframe."""
    from docx import Document
    out = tailor_docx_in_place(_v2(), source_path=FIXTURE, base="resume", out_dir=tmp_path)
    assert out.get("docx_final") == "resume_final.docx"
    final_path = tmp_path / out["docx_final"]
    assert final_path.exists()

    doc = Document(str(final_path))
    # Modified content still present in the clean version
    texts = [p.text for p in doc.paragraphs]
    assert any("Verilog testbench" in t for t in texts)
    assert any("AXI4 transaction generators" in t for t in texts)

    # Crucially: no run is colored green (RGB 0a662c).
    green_run_count = 0
    for p in doc.paragraphs:
        for r in p.runs:
            color = r.font.color.rgb
            if color is not None and str(color).lower() == "0a662c":
                green_run_count += 1
    assert green_run_count == 0, (
        f"Clean DOCX should have no green-colored runs, found {green_run_count}"
    )
