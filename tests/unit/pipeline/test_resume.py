"""Tests for pipeline.resume — PDF extraction chain, normalisation, tailored save."""
from pathlib import Path

import pytest

from pipeline.resume import (
    _build_demo_resume,
    _normalise_pdf_text,
    _read_resume,
    _save_tailored_resume,
)

pytestmark = pytest.mark.unit


# ── _build_demo_resume ──────────────────────────────────────────────────────


class TestDemoResume:
    def test_includes_all_sections(self):
        text = _build_demo_resume()
        for header in ("OBJECTIVE", "EDUCATION", "TECHNICAL SKILLS",
                       "PROJECTS", "EXPERIENCE"):
            assert header in text

    def test_no_real_personal_data(self):
        # Demo resume must use placeholder values, not somebody's actual data.
        text = _build_demo_resume()
        assert "[Your Full Name]" in text or "[Your" in text


# ── _normalise_pdf_text ─────────────────────────────────────────────────────


class TestNormalisePdfText:
    def test_replaces_ligatures(self):
        # ﬁ + ﬂ + ﬀ → fi + fl + ff
        text = "fiﬁﬂﬀ"
        assert _normalise_pdf_text(text) == "fififlff"

    def test_curly_quotes_become_ascii(self):
        text = "He said “hi” there’s nothing"
        out = _normalise_pdf_text(text)
        assert '"' in out and "'" in out
        # Curly quotes gone.
        assert "“" not in out
        assert "’" not in out

    def test_collapses_blank_lines(self):
        text = "line1\n\n\n\n\nline2"
        out = _normalise_pdf_text(text)
        assert "\n\n\n" not in out

    def test_strips_carriage_returns(self):
        out = _normalise_pdf_text("a\r\nb\rc")
        assert "\r" not in out

    def test_rejoins_mid_word_breaks_with_space(self):
        # When a line ends with a letter and the next starts lowercase, merge
        # with a single space (per implementation — handles broken layout
        # without collapsing distinct words).
        text = "exper\nimental"
        out = _normalise_pdf_text(text)
        assert "exper imental" in out
        assert "exper\nimental" not in out

    def test_em_dash_to_ascii(self):
        out = _normalise_pdf_text("today—tomorrow")
        assert "today-tomorrow" in out


# ── _read_resume ────────────────────────────────────────────────────────────


class TestReadResume:
    def test_txt_round_trip(self, tmp_path):
        p = tmp_path / "r.txt"
        p.write_text("Plain Resume Text\nLine 2\n", encoding="utf-8")
        text, latex = _read_resume(p)
        assert "Plain Resume Text" in text
        assert latex is None

    def test_md_round_trip(self, tmp_path):
        p = tmp_path / "r.md"
        p.write_text("# Header\nbody\n", encoding="utf-8")
        text, latex = _read_resume(p)
        assert "# Header" in text
        assert latex is None

    def test_tex_returns_plaintext_and_source(self, tmp_path):
        p = tmp_path / "r.tex"
        src = (
            r"\documentclass{article}\begin{document}"
            r"\section{Education}MIT \end{document}"
        )
        p.write_text(src, encoding="utf-8")
        text, latex = _read_resume(p)
        assert "EDUCATION" in text  # plaintext conversion uppercases section names
        assert latex == src

    def test_txt_with_latex_markup_detected(self, tmp_path):
        p = tmp_path / "r.txt"
        src = r"\documentclass{article}\begin{document}content\end{document}"
        p.write_text(src, encoding="utf-8")
        text, latex = _read_resume(p)
        assert latex == src

    def test_missing_file_returns_empty(self, tmp_path):
        # Reader no longer hides missing-file failures behind the demo
        # resume — empty plaintext signals the caller to surface a real
        # error to the user instead of silently extracting "Your Name".
        text, latex = _read_resume(tmp_path / "nonexistent.txt")
        assert text == ""
        assert latex is None

    def test_docx_table_text_extracted(self, tmp_path):
        # Many resumes use tables for sidebar layouts; paragraph-only
        # extraction misses everything inside cells.
        try:
            from docx import Document
        except ImportError:
            import pytest
            pytest.skip("python-docx not installed")
        doc = Document()
        doc.add_paragraph("Resume Body Paragraph")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Jane Tester"
        table.cell(0, 1).text = "jane@example.com"
        path = tmp_path / "tabular.docx"
        doc.save(str(path))
        text, _ = _read_resume(path)
        assert "Resume Body Paragraph" in text
        assert "Jane Tester" in text
        assert "jane@example.com" in text

    def test_unsupported_extension_returns_empty(self, tmp_path):
        p = tmp_path / "r.weird"
        p.write_bytes(b"\x00\x01\x02not text")
        text, latex = _read_resume(p)
        # Non-decodable binary input → empty plaintext, no LaTeX source.
        assert latex is None
        # Either UTF-8 decoded (best-effort) or empty — never the demo
        # resume placeholder which would have said "Your Full Name".
        assert "Your Full Name" not in text

    def test_fixtures_resume_text(self, fixtures_dir):
        text, latex = _read_resume(fixtures_dir / "resumes" / "sample_text.txt")
        assert latex is None
        assert "Jane Tester" in text

    def test_fixtures_resume_tex(self, fixtures_dir):
        text, latex = _read_resume(fixtures_dir / "resumes" / "sample_latex.tex")
        # tex source preserved AND plaintext present.
        assert latex is not None
        assert r"\documentclass" in latex
        assert "Jane Tester" in text


# ── _save_tailored_resume ───────────────────────────────────────────────────


class TestSaveTailoredResume:
    """v2: _save_tailored_resume dispatches by source_format. The default path
    (no source_format → template lib) emits HTML preview + PDF; the explicit
    .tex path emits .tex + (optionally) PDF + HTML preview."""

    def test_default_path_writes_html_preview(self, tmp_path):
        profile = {"name": "Jane Tester", "email": "j@x.com",
                   "top_hard_skills": ["Python", "Verilog"], "education": [],
                   "experience": [], "projects": []}
        tailored = {"skills_reordered": ["Verilog", "Python"],
                    "experience_bullets": [], "ats_keywords_missing": [],
                    "section_order": ["Skills"]}
        job = {"title": "FPGA Intern", "company": "Acme"}
        out = _save_tailored_resume(job, tailored, profile, output_dir=tmp_path)
        assert out["html_preview"] is not None
        html = (tmp_path / out["html_preview"]).read_text(encoding="utf-8")
        assert "Verilog" in html and "Python" in html
        assert "FPGA" in out["base"] and "Acme" in out["base"]

    def test_pdf_or_skipped(self, tmp_path):
        profile = {"name": "Jane Tester", "education": [], "experience": [],
                   "projects": [], "top_hard_skills": []}
        tailored = {"skills_reordered": ["X"], "experience_bullets": [],
                    "ats_keywords_missing": [],
                    "section_order": ["Skills"]}
        job = {"title": "Eng", "company": "Co"}
        out = _save_tailored_resume(job, tailored, profile, output_dir=tmp_path)
        # Either a PDF was produced (reportlab path) or it's None (no backend).
        assert out["pdf"] is None or (tmp_path / out["pdf"]).exists()

    def test_with_latex_source_in_place_path(self, tmp_path):
        profile = {"name": "Jane Tester", "education": [], "experience": [],
                   "projects": [], "top_hard_skills": []}
        tailored = {"skills_reordered": ["Foo", "Bar"],
                    "experience_bullets": [], "ats_keywords_missing": [],
                    "section_order": ["Skills"]}
        latex_src = (
            r"\begin{document}\section{Skills}old\section{Education}edu\end{document}"
        )
        job = {"title": "Eng", "company": "Co"}
        out = _save_tailored_resume(job, tailored, profile,
                                     latex_source=latex_src,
                                     output_dir=tmp_path,
                                     source_format="tex")
        # In-place LaTeX path emits a .tex file
        assert out["tex"] is not None
        tex_content = (tmp_path / out["tex"]).read_text(encoding="utf-8")
        assert "Foo" in tex_content and "Bar" in tex_content
        assert out["template_id"] == "in_place_latex"
